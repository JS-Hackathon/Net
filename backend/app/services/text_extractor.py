import io
import logging
from typing import Optional

logger = logging.getLogger(__name__)

class TextExtractor:
    @staticmethod
    def extract_text_from_pdf(file_bytes: bytes) -> str:
        """Extract text from PDF file bytes using pypdf."""
        try:
            import pypdf
            pdf_file = io.BytesIO(file_bytes)
            reader = pypdf.PdfReader(pdf_file)
            text = ""
            for page_num, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {e}")
            raise ValueError(f"Không thể trích xuất văn bản từ file PDF: {str(e)}")

    @staticmethod
    def extract_text_from_docx(file_bytes: bytes) -> str:
        """Extract text from DOCX file bytes using python-docx."""
        try:
            import docx
            docx_file = io.BytesIO(file_bytes)
            doc = docx.Document(docx_file)
            text_parts = []
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract table cells text
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        text_parts.append(" | ".join(row_text))

            return "\n".join(text_parts).strip()
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {e}")
            raise ValueError(f"Không thể trích xuất văn bản từ file DOCX: {str(e)}")

    @classmethod
    def extract_text(cls, file_bytes: bytes, file_type: str) -> str:
        """Helper routing to extract text based on file format (pdf or docx)."""
        ftype = file_type.lower().strip(".")
        if ftype == "pdf":
            return cls.extract_text_from_pdf(file_bytes)
        elif ftype in ("docx", "doc"):
            return cls.extract_text_from_docx(file_bytes)
        else:
            raise ValueError(f"Định dạng file không được hỗ trợ để trích xuất văn bản: {file_type}")
